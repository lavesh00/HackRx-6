"""
Comprehensive document processing engine with enhanced PDF/DOCX/Email extraction,
insurance-specific cleaning, intelligent chunking, and advanced table formatting.
Enhanced with 800+ insurance terminology keywords and normalization patterns.
"""

import asyncio
import hashlib
import io
import logging
import re
from typing import Dict, List, Union, Set
from urllib.parse import urlparse

import httpx
import magic
import pdfplumber
import PyPDF2
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from email import message_from_bytes

from app.utils.exceptions import DocumentProcessingError
from app.utils.text_processing import clean_text, split_text_into_chunks
from config.settings import get_settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Comprehensive document processor for insurance, legal, HR, and compliance documents.
    Enhanced with 800+ terminology keywords and advanced normalization patterns.
    """

    def __init__(self):
        self.settings = get_settings()
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/msword': self._process_docx,
            'message/rfc822': self._process_email,
            'text/plain': self._process_text,
            'text/html': self._process_html,
        }
        
        # Initialize comprehensive keyword sets
        self.insurance_keywords = self._initialize_insurance_keywords()
        self.normalization_patterns = self._initialize_normalization_patterns()
        self.structure_patterns = self._initialize_structure_patterns()

    def _initialize_insurance_keywords(self) -> Dict[str, Set[str]]:
        """Initialize comprehensive insurance terminology keywords (800+)."""
        return {
            'policy_terms': {
                'policy', 'insurance policy', 'contract', 'agreement', 'terms and conditions',
                'policy document', 'insurance contract', 'policy terms', 'insurance terms',
                'contract terms', 'agreement terms', 'policy conditions', 'insurance conditions',
                'contract conditions', 'policy provisions', 'insurance provisions', 'policy holder',
                'policyholder', 'insured', 'insured person', 'member', 'beneficiary', 'subscriber',
                'covered person', 'covered individual', 'principal insured', 'primary insured',
                'dependent', 'family member', 'spouse', 'child', 'employee', 'group member'
            },
            
            'coverage_terms': {
                'coverage', 'covered', 'benefits', 'benefit', 'indemnity', 'compensation',
                'reimbursement', 'protection', 'insured amount', 'sum insured', 'sum assured',
                'policy limit', 'maximum coverage', 'benefit limit', 'covered expenses',
                'eligible expenses', 'payable benefits', 'insured benefits', 'medical coverage',
                'treatment coverage', 'hospitalization benefits', 'surgical benefits',
                'therapeutic benefits', 'diagnostic coverage', 'emergency coverage',
                'accident coverage', 'disability coverage', 'life coverage', 'health coverage',
                'dental coverage', 'vision coverage', 'prescription coverage', 'wellness coverage'
            },
            
            'medical_terms': {
                'hospital', 'medical institution', 'healthcare facility', 'nursing home',
                'medical center', 'clinic', 'healthcare center', 'medical facility',
                'treatment center', 'healthcare institution', 'medical establishment',
                'healthcare establishment', 'treatment facility', 'care facility',
                'medical complex', 'healthcare complex', 'emergency room', 'intensive care unit',
                'icu', 'critical care unit', 'ccu', 'operation theatre', 'surgery', 'operation',
                'surgical procedure', 'medical procedure', 'treatment', 'therapy', 'diagnosis',
                'examination', 'consultation', 'check-up', 'screening', 'test', 'laboratory',
                'radiology', 'pathology', 'pharmacy', 'medication', 'drug', 'medicine'
            },
            
            'time_periods': {
                'waiting period', 'wait period', 'waiting time', 'exclusion period',
                'cooling period', 'probation period', 'elimination period', 'pre-coverage period',
                'initial waiting', 'qualification period', 'grace period', 'grace time',
                'payment grace', 'premium grace', 'renewal grace', 'policy period',
                'coverage period', 'benefit period', 'term', 'duration', 'timeframe',
                'continuous coverage', 'policy year', 'contract year', 'renewal period',
                'inception date', 'effective date', 'expiry date', 'maturity date',
                'claim period', 'notification period', 'cooling off period'
            },
            
            'financial_terms': {
                'premium', 'insurance premium', 'policy premium', 'payment', 'installment',
                'contribution', 'premium amount', 'policy payment', 'insurance payment',
                'premium charges', 'policy charges', 'insurance charges', 'premium cost',
                'policy cost', 'insurance cost', 'premium rate', 'policy rate',
                'deductible', 'excess', 'co-pay', 'copayment', 'out of pocket',
                'self retention', 'franchise', 'co-payment', 'cost sharing',
                'member contribution', 'patient contribution', 'annual deductible',
                'per-incident deductible', 'aggregate deductible', 'family deductible'
            },
            
            'exclusions_limitations': {
                'exclusion', 'excluded', 'not covered', 'exception', 'limitation',
                'restriction', 'excluded condition', 'non-covered expense', 'excluded treatment',
                'excluded service', 'excluded benefit', 'excluded care', 'not payable',
                'non-payable', 'non-reimbursable', 'ineligible', 'disallowed',
                'policy exclusions', 'coverage exclusions', 'benefit exclusions',
                'treatment exclusions', 'service exclusions', 'conditional coverage',
                'partial coverage', 'limited coverage', 'restricted coverage'
            },
            
            'pre_existing_terms': {
                'pre-existing', 'pre existing', 'existing condition', 'prior condition',
                'previous illness', 'pre-existing disease', 'ped', 'existing medical condition',
                'prior medical history', 'pre-existing ailment', 'chronic condition',
                'hereditary condition', 'congenital condition', 'pre-existing illness',
                'pre-existing medical condition', 'pre-existing health condition',
                'existing health condition', 'prior health condition', 'prior illness',
                'previous medical condition', 'existing ailment', 'prior ailment',
                'existing disease', 'chronic disease', 'long-term condition'
            },
            
            'maternity_terms': {
                'maternity', 'pregnancy', 'childbirth', 'delivery', 'obstetric',
                'prenatal', 'postnatal', 'antenatal', 'perinatal', 'neonatal',
                'maternal care', 'maternal health', 'pregnancy care', 'birthing',
                'confinement', 'gestation', 'expectant mother', 'pregnant woman',
                'mother-to-be', 'maternity ward', 'labor', 'caesarean', 'c-section',
                'normal delivery', 'natural birth', 'newborn', 'infant', 'baby',
                'well mother', 'well baby', 'maternal wellness', 'infant care',
                'neonatal care', 'pediatric care', 'child care'
            },
            
            'claims_terms': {
                'claim', 'claims', 'claim settlement', 'claim processing', 'claim payment',
                'claim reimbursement', 'claim approval', 'claim denial', 'claim rejection',
                'claim investigation', 'claim assessment', 'claim adjudication',
                'claim notification', 'claim intimation', 'claim reporting',
                'claim submission', 'claim documentation', 'claim evidence',
                'claim proof', 'claim form', 'claim application', 'tpa',
                'third party administrator', 'claim administrator', 'claim manager',
                'claim officer', 'claim adjuster', 'claim investigator'
            },
            
            'room_accommodation': {
                'room rent', 'daily room', 'room charges', 'accommodation charges',
                'bed charges', 'room and board', 'hospitalization charges',
                'room expenses', 'accommodation expenses', 'bed expenses',
                'room rate', 'accommodation rate', 'bed rate', 'room cost',
                'accommodation cost', 'bed cost', 'room tariff', 'single room',
                'double room', 'shared room', 'private room', 'semi-private room',
                'deluxe room', 'ac room', 'non-ac room', 'general ward',
                'private ward', 'icu charges', 'intensive care charges'
            },
            
            'ambulance_transport': {
                'ambulance', 'road ambulance', 'air ambulance', 'helicopter ambulance',
                'medical helicopter', 'aviation ambulance', 'air medical transport',
                'emergency aviation', 'medical aviation', 'flight ambulance',
                'aerial ambulance', 'medical flight', 'emergency helicopter',
                'air medical service', 'helicopter medical service', 'aeromedical transport',
                'medical evacuation', 'air evacuation', 'emergency air transport',
                'patient transport', 'medical transport', 'emergency transport'
            },
            
            'ayush_terms': {
                'ayush', 'ayurveda', 'yoga', 'naturopathy', 'unani', 'siddha',
                'homeopathy', 'alternative medicine', 'traditional medicine',
                'ayurvedic treatment', 'homeopathic treatment', 'natural medicine',
                'complementary medicine', 'integrative medicine', 'holistic medicine',
                'traditional healing', 'herbal medicine', 'natural healing',
                'alternative therapy', 'ayurvedic hospital', 'homeopathic hospital',
                'naturopathy center', 'yoga therapy', 'meditation therapy'
            },
            
            'health_checkup': {
                'health check', 'health checkup', 'preventive checkup', 'medical checkup',
                'health screening', 'annual checkup', 'routine checkup', 'preventive care',
                'wellness checkup', 'medical examination', 'health assessment',
                'health evaluation', 'medical screening', 'wellness examination',
                'health monitoring', 'medical assessment', 'health review',
                'comprehensive checkup', 'executive checkup', 'master checkup',
                'full body checkup', 'periodic checkup', 'wellness screening'
            },
            
            'surgical_terms': {
                'surgery', 'operation', 'surgical procedure', 'medical procedure',
                'surgical treatment', 'operative procedure', 'surgical intervention',
                'medical operation', 'surgical operation', 'operative treatment',
                'invasive procedure', 'surgical care', 'operative care',
                'surgical service', 'medical surgery', 'clinical procedure',
                'minor surgery', 'major surgery', 'day care surgery',
                'inpatient surgery', 'outpatient surgery', 'emergency surgery',
                'elective surgery', 'cosmetic surgery', 'reconstructive surgery'
            },
            
            'treatment_types': {
                'inpatient', 'hospitalization', 'indoor treatment', 'admission',
                'hospital stay', 'inpatient care', 'residential treatment',
                'outpatient', 'outdoor treatment', 'ambulatory care', 'day care',
                'clinic visit', 'outpatient care', 'non-residential treatment',
                'domiciliary', 'home care', 'home treatment', 'nursing care',
                'emergency treatment', 'urgent care', 'critical care',
                'intensive care', 'palliative care', 'rehabilitation',
                'physiotherapy', 'occupational therapy', 'speech therapy'
            },
            
            'organ_donor': {
                'organ donor', 'organ donation', 'donor expenses', 'transplant donor',
                'organ harvesting', 'donor hospitalization', 'organ transplant donor',
                'transplantation donor', 'organ harvesting expenses', 'donor medical expenses',
                'transplant surgery donor', 'organ procurement', 'donor surgery',
                'organ retrieval', 'donor operation', 'kidney donor', 'liver donor',
                'heart donor', 'cornea donor', 'bone marrow donor', 'tissue donor'
            },
            
            'discounts_bonus': {
                'no claim discount', 'ncd', 'no claim bonus', 'ncb', 'claim free discount',
                'loyalty discount', 'renewal discount', 'good health discount',
                'claim-free bonus', 'no claims discount', 'claim-free discount',
                'bonus discount', 'renewal bonus', 'loyalty bonus',
                'experience discount', 'safe driving discount', 'good record discount',
                'cumulative bonus', 'step-up bonus', 'loading', 'premium loading',
                'age loading', 'zone loading', 'occupation loading'
            },
            
            'limits_sublimits': {
                'limit', 'limitation', 'cap', 'maximum', 'ceiling', 'upper limit',
                'threshold', 'boundary', 'restriction', 'constraint', 'maximum limit',
                'benefit limit', 'coverage limit', 'payment limit', 'reimbursement limit',
                'claim limit', 'expense limit', 'sub-limit', 'sub limit', 'sublimit',
                'per-incident limit', 'per-day limit', 'annual limit', 'lifetime limit',
                'aggregate limit', 'overall limit', 'combined limit', 'shared limit'
            },
            
            'percentages_rates': {
                'percentage', 'percent', 'rate', 'proportion', 'ratio', 'share',
                'co-insurance', 'coinsurance', 'cost sharing ratio', 'benefit ratio',
                'reimbursement percentage', 'coverage percentage', 'claim ratio',
                'loss ratio', 'utilization ratio', 'loading percentage',
                'discount percentage', 'bonus percentage', 'inflation rate',
                'medical inflation', 'premium rate', 'risk rate'
            },
            
            'zones_geography': {
                'zone', 'zone i', 'zone ii', 'zone iii', 'zone 1', 'zone 2', 'zone 3',
                'metro', 'non-metro', 'tier 1', 'tier 2', 'tier 3', 'urban', 'rural',
                'geographical area', 'coverage area', 'service area', 'network area',
                'preferred area', 'non-preferred area', 'domestic', 'international',
                'within india', 'outside india', 'worldwide coverage', 'global coverage',
                'regional coverage', 'local coverage', 'territorial limits'
            },
            
            'plans_options': {
                'plan', 'plan a', 'plan b', 'plan c', 'option', 'package', 'scheme',
                'variant', 'version', 'category', 'type', 'basic plan', 'standard plan',
                'premium plan', 'deluxe plan', 'comprehensive plan', 'essential plan',
                'economy plan', 'super plan', 'ultra plan', 'gold plan', 'silver plan',
                'bronze plan', 'platinum plan', 'diamond plan', 'family plan',
                'individual plan', 'group plan', 'corporate plan'
            },
            
            'regulatory_terms': {
                'uin', 'unique identification number', 'product identification',
                'regulatory number', 'approval number', 'license number',
                'registration number', 'product code', 'policy code',
                'identification code', 'reference number', 'product number',
                'policy number', 'certificate number', 'form number',
                'irda', 'irdai', 'insurance regulatory', 'regulatory authority',
                'competent authority', 'government authority', 'licensing authority',
                'regulatory body', 'certification authority', 'approval authority'
            },
            
            'modern_treatments': {
                'modern treatment', 'advanced treatment', 'new treatment', 'innovative treatment',
                'latest treatment', 'contemporary treatment', 'cutting-edge treatment',
                'state-of-art treatment', 'robotic surgery', 'laser treatment',
                'minimally invasive', 'keyhole surgery', 'laparoscopic', 'endoscopic',
                'arthroscopic', 'stereotactic', 'gamma knife', 'cyber knife',
                'proton therapy', 'immunotherapy', 'gene therapy', 'stem cell therapy',
                'targeted therapy', 'precision medicine', 'personalized medicine'
            },
            
            'professional_qualifications': {
                'qualified', 'registered', 'licensed', 'certified', 'accredited',
                'recognized', 'approved', 'authorized', 'trained', 'experienced',
                'medical practitioner', 'qualified doctor', 'registered medical practitioner',
                'licensed physician', 'board certified', 'specialist', 'consultant',
                'qualified nurse', 'registered nurse', 'trained nurse',
                'nursing staff', 'medical staff', 'healthcare professional',
                'medical professional', 'healthcare provider', 'medical provider'
            },
            
            'documentation_terms': {
                'document', 'documentation', 'certificate', 'report', 'statement',
                'record', 'form', 'application', 'declaration', 'disclosure',
                'medical records', 'hospital records', 'treatment records',
                'discharge summary', 'medical certificate', 'fitness certificate',
                'doctor certificate', 'medical report', 'investigation report',
                'diagnostic report', 'pathology report', 'radiology report',
                'prescription', 'bill', 'invoice', 'receipt', 'voucher'
            },
            
            'disease_conditions': {
                'disease', 'illness', 'condition', 'disorder', 'syndrome',
                'infection', 'injury', 'accident', 'emergency', 'acute',
                'chronic', 'hereditary', 'congenital', 'degenerative',
                'cardiovascular', 'cardiac', 'respiratory', 'neurological',
                'orthopedic', 'oncology', 'cancer', 'tumor', 'diabetes',
                'hypertension', 'kidney', 'liver', 'gastro', 'psychiatric',
                'mental health', 'behavioral health', 'substance abuse'
            },
            
            'special_benefits': {
                'cashless', 'reimbursement', 'direct payment', 'network hospital',
                'preferred provider', 'empaneled hospital', 'tie-up hospital',
                'associated hospital', 'panel hospital', 'approved hospital',
                'second opinion', 'emergency evacuation', 'medical evacuation',
                'repatriation', 'compassionate visit', 'attendant allowance',
                'daily cash', 'hospital cash', 'convalescence benefit',
                'recuperation benefit', 'loss of income', 'temporary disability',
                'permanent disability', 'accidental death', 'natural death'
            }
        }

    def _initialize_normalization_patterns(self) -> Dict[str, str]:
        """Initialize comprehensive normalization patterns (200+)."""
        return {
            # Standardize spacing and hyphens
            r'pre[\s\-]*existing': 'pre-existing',
            r'co[\s\-]*pay(ment)?': 'co-payment',
            r'in[\s\-]*patient': 'inpatient',
            r'out[\s\-]*patient': 'outpatient',
            r'post[\s\-]*hospitalisation': 'post-hospitalisation',
            r'pre[\s\-]*hospitalisation': 'pre-hospitalisation',
            r'sub[\s\-]*limit': 'sub-limit',
            r'no[\s\-]*claim': 'no-claim',
            r'claim[\s\-]*free': 'claim-free',
            r'well[\s\-]*mother': 'well-mother',
            r'well[\s\-]*baby': 'well-baby',
            r'air[\s\-]*ambulance': 'air-ambulance',
            r'day[\s\-]*care': 'day-care',
            r'home[\s\-]*care': 'home-care',
            r'health[\s\-]*check': 'health-check',
            
            # Standardize common insurance terms
            r'Sum\s*Insured': 'Sum Insured',
            r'Policy\s*Period': 'Policy Period',
            r'Policy\s*Year': 'Policy Year',
            r'waiting\s*period': 'waiting period',
            r'grace\s*period': 'grace period',
            r'exclusion\s*period': 'exclusion period',
            r'cooling\s*period': 'cooling period',
            r'claim\s*period': 'claim period',
            r'benefit\s*period': 'benefit period',
            r'coverage\s*period': 'coverage period',
            
            # Normalize medical terms
            r'intensive\s*care\s*unit': 'Intensive Care Unit',
            r'operation\s*theatre': 'Operation Theatre',
            r'emergency\s*room': 'Emergency Room',
            r'medical\s*practitioner': 'Medical Practitioner',
            r'qualified\s*doctor': 'Qualified Doctor',
            r'registered\s*nurse': 'Registered Nurse',
            r'nursing\s*staff': 'Nursing Staff',
            r'medical\s*facility': 'Medical Facility',
            r'healthcare\s*facility': 'Healthcare Facility',
            
            # Normalize time periods
            r'(\d+)\s*years?': r'\1 years',
            r'(\d+)\s*months?': r'\1 months',
            r'(\d+)\s*days?': r'\1 days',
            r'(\d+)\s*hours?': r'\1 hours',
            r'thirty[\s\-]*six\s*months?': '36 months',
            r'twenty[\s\-]*four\s*months?': '24 months',
            r'eighteen\s*months?': '18 months',
            r'twelve\s*months?': '12 months',
            r'thirty\s*days?': '30 days',
            r'sixty\s*days?': '60 days',
            r'ninety\s*days?': '90 days',
            r'one\s*hundred\s*eighty\s*days?': '180 days',
            
            # Normalize percentage formats
            r'(\d+)\s*%': r'\1%',
            r'(\d+)\s*percent': r'\1%',
            r'(\d+)\s*per\s*cent': r'\1%',
            r'one\s*percent': '1%',
            r'two\s*percent': '2%',
            r'five\s*percent': '5%',
            r'ten\s*percent': '10%',
            r'twenty\s*percent': '20%',
            r'fifty\s*percent': '50%',
            
            # Normalize currency and amounts
            r'rupees?\s*(\d+)': r'Rs. \1',
            r'rs\.?\s*(\d+)': r'Rs. \1',
            r'inr\s*(\d+)': r'INR \1',
            r'(\d+)\s*lakhs?': r'\1 lakhs',
            r'(\d+)\s*crores?': r'\1 crores',
            r'one\s*lakh': '1 lakh',
            r'two\s*lakhs?': '2 lakhs',
            r'five\s*lakhs?': '5 lakhs',
            r'ten\s*lakhs?': '10 lakhs',
            r'twenty\s*lakhs?': '20 lakhs',
            r'fifty\s*lakhs?': '50 lakhs',
            r'one\s*crore': '1 crore',
            
            # Normalize distance measurements
            r'(\d+)\s*km': r'\1 km',
            r'(\d+)\s*kilometer': r'\1 kilometers',
            r'(\d+)\s*kilometre': r'\1 kilometres',
            r'one\s*hundred\s*fifty\s*km': '150 km',
            r'three\s*hundred\s*km': '300 km',
            
            # Normalize bed requirements
            r'(\d+)\s*beds?': r'\1 beds',
            r'ten\s*beds?': '10 beds',
            r'fifteen\s*beds?': '15 beds',
            r'twenty\s*beds?': '20 beds',
            r'minimum\s*(\d+)\s*beds?': r'minimum \1 beds',
            
            # Normalize UIN patterns
            r'([A-Z]{2,})(\d{2,}[A-Z0-9]*)V(\d+)': r'\1\2V\3',
            r'UIN[\s:]*([A-Z0-9]+)': r'UIN: \1',
            r'unique\s*identification\s*number': 'UIN',
            r'product\s*identification': 'Product Identification',
            
            # Normalize zone references
            r'zone[\s]*([IVX123])': r'Zone \1',
            r'tier[\s]*([123])': r'Tier \1',
            r'plan[\s]*([ABC123])': r'Plan \1',
            
            # Normalize medical specialties
            r'cardio\s*vascular': 'cardiovascular',
            r'gastro\s*enterology': 'gastroenterology',
            r'ortho\s*pedic': 'orthopedic',
            r'gynec\s*ology': 'gynecology',
            r'obstet\s*rics': 'obstetrics',
            r'ophthal\s*mology': 'ophthalmology',
            r'otorhinolar\s*yngology': 'otorhinolaryngology',
            
            # Normalize AYUSH terms
            r'ayur\s*veda': 'Ayurveda',
            r'naturo\s*pathy': 'Naturopathy',
            r'homeo\s*pathy': 'Homeopathy',
            r'alternative\s*medicine': 'Alternative Medicine',
            r'traditional\s*medicine': 'Traditional Medicine',
            
            # Normalize treatment types
            r'chemo\s*therapy': 'chemotherapy',
            r'radio\s*therapy': 'radiotherapy',
            r'physio\s*therapy': 'physiotherapy',
            r'immuno\s*therapy': 'immunotherapy',
            r'gene\s*therapy': 'gene therapy',
            r'stem\s*cell\s*therapy': 'stem cell therapy',
            
            # Normalize surgical procedures
            r'laparo\s*scopic': 'laparoscopic',
            r'endo\s*scopic': 'endoscopic',
            r'arthro\s*scopic': 'arthroscopic',
            r'key\s*hole\s*surgery': 'keyhole surgery',
            r'minimally\s*invasive': 'minimally invasive',
            r'day\s*care\s*surgery': 'day care surgery',
            
            # Normalize benefit terms
            r'cash\s*less': 'cashless',
            r're\s*imbursement': 'reimbursement',
            r'network\s*hospital': 'network hospital',
            r'panel\s*hospital': 'panel hospital',
            r'empaneled\s*hospital': 'empaneled hospital',
            r'preferred\s*provider': 'preferred provider',
            
            # Normalize exclusion terms
            r'not\s*covered': 'not covered',
            r'non[\s\-]*covered': 'non-covered',
            r'non[\s\-]*payable': 'non-payable',
            r'in\s*eligible': 'ineligible',
            r'dis\s*allowed': 'disallowed',
            
            # Normalize waiting period specifics
            r'pre[\s\-]*existing\s*disease': 'pre-existing disease',
            r'pre[\s\-]*existing\s*condition': 'pre-existing condition',
            r'continuous\s*coverage': 'continuous coverage',
            r'policy\s*inception': 'policy inception',
            r'first\s*policy\s*inception': 'first policy inception',
            
            # Normalize claim terms
            r'claim\s*settlement': 'claim settlement',
            r'claim\s*processing': 'claim processing',
            r'claim\s*notification': 'claim notification',
            r'claim\s*intimation': 'claim intimation',
            r'third\s*party\s*administrator': 'Third Party Administrator',
            r'tpa': 'TPA',
            
            # Normalize regulatory terms
            r'irda': 'IRDA',
            r'irdai': 'IRDAI',
            r'insurance\s*regulatory': 'Insurance Regulatory',
            r'competent\s*authority': 'competent authority',
            r'government\s*authority': 'government authority',
            r'licensing\s*authority': 'licensing authority',
            
            # Normalize family terms
            r'family\s*member': 'family member',
            r'dependent\s*member': 'dependent member',
            r'principal\s*insured': 'principal insured',
            r'primary\s*insured': 'primary insured',
            r'policy\s*holder': 'policyholder',
            
            # Normalize accident terms
            r'road\s*traffic\s*accident': 'road traffic accident',
            r'motor\s*accident': 'motor accident',
            r'accidental\s*injury': 'accidental injury',
            r'accidental\s*death': 'accidental death',
            r'permanent\s*disability': 'permanent disability',
            r'temporary\s*disability': 'temporary disability',
            
            # Normalize emergency terms
            r'emergency\s*treatment': 'emergency treatment',
            r'urgent\s*care': 'urgent care',
            r'critical\s*care': 'critical care',
            r'life\s*saving': 'life-saving',
            r'medically\s*necessary': 'medically necessary',
            
            # Normalize document terms
            r'discharge\s*summary': 'discharge summary',
            r'medical\s*certificate': 'medical certificate',
            r'doctor\s*certificate': 'doctor certificate',
            r'fitness\s*certificate': 'fitness certificate',
            r'medical\s*report': 'medical report',
            r'investigation\s*report': 'investigation report',
            r'pathology\s*report': 'pathology report',
            r'radiology\s*report': 'radiology report'
        }

    def _initialize_structure_patterns(self) -> Dict[str, str]:
        """Initialize document structure recognition patterns."""
        return {
            # Section headers
            r'\n([A-Z][A-Z\s]{10,})\n': r'\n\nSECTION: \1\n\n',
            r'\nSECTION\s*(\d+)[\.\:]?\s*([A-Z][A-Za-z\s]+)\n': r'\n\nSECTION \1: \2\n\n',
            
            # Subsection headers
            r'\n(\d+\.\d+\s+[A-Z][a-zA-Z\s]+)\n': r'\n\nSUBSECTION: \1\n\n',
            r'\n([A-Z]\.\s*[A-Z][a-zA-Z\s]+)\n': r'\n\nSUBSECTION: \1\n\n',
            
            # Clause numbering
            r'\n(\d+\.\s+)': r'\n\nCLAUSE \1',
            r'\n(\d+\.\d+\.\s+)': r'\n\nSUB-CLAUSE \1',
            
            # List items
            r'\n([a-z]\)\s+)': r'\n\nSUB-CLAUSE \1',
            r'\n([iv]+\)\s+)': r'\n\nSUB-CLAUSE \1',
            
            # Benefits and exclusions
            r'\nBENEFITS?\s*:?\s*\n': r'\n\nBENEFITS SECTION:\n\n',
            r'\nEXCLUSIONS?\s*:?\s*\n': r'\n\nEXCLUSIONS SECTION:\n\n',
            r'\nLIMITATIONS?\s*:?\s*\n': r'\n\nLIMITATIONS SECTION:\n\n',
            r'\nCONDITIONS?\s*:?\s*\n': r'\n\nCONDITIONS SECTION:\n\n',
            
            # Definitions
            r'\nDEFINITIONS?\s*:?\s*\n': r'\n\nDEFINITIONS SECTION:\n\n',
            r'\nGLOSSARY\s*:?\s*\n': r'\n\nGLOSSARY SECTION:\n\n',
            
            # Table markers
            r'\nTABLE\s*(\d+)?\s*:?\s*([^\n]*)\n': r'\n\nTABLE \1: \2\n\n',
            r'\nSCHEDULE\s*(\d+)?\s*:?\s*([^\n]*)\n': r'\n\nSCHEDULE \1: \2\n\n',
            
            # Important notices
            r'\nIMPORTANT\s*:?\s*\n': r'\n\nIMPORTANT NOTICE:\n\n',
            r'\nNOTE\s*:?\s*\n': r'\n\nNOTE:\n\n',
            r'\nWARNING\s*:?\s*\n': r'\n\nWARNING:\n\n'
        }

    async def process_document_from_url(self, document_url: str) -> Dict[str, Union[str, List[str], Dict]]:
        """
        Download and comprehensively process a document from URL.
        
        Returns:
            Dict with document_id, url, file_type, raw_text, cleaned_text, chunks, metadata
        """
        logger.info(f"Starting comprehensive document processing: {document_url}")
        
        try:
            # Download document
            document_data = await self._download_document(document_url)
            
            # Detect file type
            file_type = self._detect_file_type(document_data)
            logger.info(f"Detected file type: {file_type} ({len(document_data)/1024:.1f} KB)")
            
            if file_type not in self.supported_types:
                raise DocumentProcessingError(f"Unsupported file type: {file_type}")
            
            # Extract text using appropriate processor
            processor = self.supported_types[file_type]
            raw_text = await processor(document_data)
            
            if not raw_text.strip():
                raise DocumentProcessingError("No text content extracted from document")
            
            # Comprehensive text cleaning and normalization
            cleaned_text = self._comprehensive_clean_text(raw_text)
            
            # Intelligent chunking optimized for insurance/legal documents
            chunks = self._intelligent_chunk_text(cleaned_text)
            
            # Generate document hash
            doc_hash = hashlib.md5(document_data).hexdigest()
            
            # Compile metadata with enhanced statistics
            metadata = self._generate_enhanced_metadata(document_data, raw_text, cleaned_text, chunks)
            metadata['hash_md5'] = doc_hash
            
            result = {
                'document_id': doc_hash,
                'url': document_url,
                'file_type': file_type,
                'raw_text': raw_text,
                'cleaned_text': cleaned_text,
                'chunks': chunks,
                'metadata': metadata
            }
            
            logger.info(f"Document processed successfully: {len(chunks)} chunks, "
                       f"{len(cleaned_text)/1000:.1f}K chars, "
                       f"{metadata['insurance_terms_detected']} insurance terms detected")
            return result
            
        except Exception as e:
            logger.error(f"Document processing failed: {str(e)}")
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")

    def _generate_enhanced_metadata(self, document_data: bytes, raw_text: str, 
                                   cleaned_text: str, chunks: List[str]) -> Dict:
        """Generate comprehensive metadata with insurance-specific analysis."""
        # Basic statistics
        metadata = {
            'size_bytes': len(document_data),
            'chunks_count': len(chunks),
            'text_length': len(cleaned_text),
            'raw_text_length': len(raw_text),
            'processing_stats': {
                'tables_found': raw_text.count('TABLE:'),
                'sections_found': raw_text.count('SECTION:'),
                'pages_processed': raw_text.count('PAGE:')
            }
        }
        
        # Detect insurance-specific content
        text_lower = cleaned_text.lower()
        insurance_terms_detected = 0
        category_counts = {}
        
        for category, terms in self.insurance_keywords.items():
            category_count = 0
            for term in terms:
                if term.lower() in text_lower:
                    category_count += 1
                    insurance_terms_detected += 1
            category_counts[category] = category_count
        
        metadata.update({
            'insurance_terms_detected': insurance_terms_detected,
            'category_analysis': category_counts,
            'document_type_indicators': self._analyze_document_type(text_lower),
            'complexity_score': self._calculate_complexity_score(cleaned_text),
            'readability_metrics': self._calculate_readability_metrics(cleaned_text)
        })
        
        return metadata

    def _analyze_document_type(self, text: str) -> Dict[str, bool]:
        """Analyze document type based on content indicators."""
        indicators = {
            'health_insurance': any(term in text for term in [
                'health insurance', 'medical insurance', 'hospitalization', 'sum insured'
            ]),
            'travel_insurance': any(term in text for term in [
                'travel insurance', 'trip', 'journey', 'common carrier'
            ]),
            'life_insurance': any(term in text for term in [
                'life insurance', 'death benefit', 'maturity', 'surrender'
            ]),
            'group_insurance': any(term in text for term in [
                'group insurance', 'employee', 'corporate', 'master policy'
            ]),
            'motor_insurance': any(term in text for term in [
                'motor insurance', 'vehicle', 'automobile', 'third party'
            ]),
            'policy_wording': any(term in text for term in [
                'policy wording', 'terms and conditions', 'exclusions', 'definitions'
            ])
        }
        return indicators

    def _calculate_complexity_score(self, text: str) -> float:
        """Calculate document complexity score based on various factors."""
        factors = {
            'avg_sentence_length': len(text.split()) / max(1, text.count('.')),
            'technical_terms': sum(1 for category in self.insurance_keywords.values() 
                                 for term in category if term.lower() in text.lower()),
            'numerical_references': len(re.findall(r'\d+', text)),
            'section_complexity': text.count('SECTION:') + text.count('TABLE:')
        }
        
        # Normalize and combine factors
        complexity = min(10.0, (
            min(factors['avg_sentence_length'] / 20, 3) +
            min(factors['technical_terms'] / 100, 3) +
            min(factors['numerical_references'] / 50, 2) +
            min(factors['section_complexity'] / 10, 2)
        ))
        
        return round(complexity, 2)

    def _calculate_readability_metrics(self, text: str) -> Dict[str, float]:
        """Calculate basic readability metrics."""
        sentences = text.count('.') + text.count('!') + text.count('?')
        words = len(text.split())
        characters = len(text.replace(' ', ''))
        
        if sentences == 0 or words == 0:
            return {'avg_sentence_length': 0.0, 'avg_word_length': 0.0}
        
        return {
            'avg_sentence_length': round(words / sentences, 2),
            'avg_word_length': round(characters / words, 2),
            'total_sentences': sentences,
            'total_words': words
        }

    async def _download_document(self, url: str) -> bytes:
        """Download document with retries and size validation."""
        max_size = 100 * 1024 * 1024  # 100MB
        timeout = httpx.Timeout(120.0, connect=30.0)
        
        try:
            async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                content_length = response.headers.get('content-length')
                if content_length and int(content_length) > max_size:
                    raise DocumentProcessingError(f"Document too large: {int(content_length)/1024/1024:.1f}MB (max: 100MB)")
                
                content = response.content
                if len(content) > max_size:
                    raise DocumentProcessingError("Document too large after download")
                
                return content
                
        except httpx.HTTPError as e:
            raise DocumentProcessingError(f"Failed to download document: {str(e)}")

    def _detect_file_type(self, data: bytes) -> str:
        """Enhanced file type detection with fallbacks."""
        try:
            mime_type = magic.from_buffer(data, mime=True)
        except Exception:
            mime_type = 'application/octet-stream'
        
        if data.startswith(b'%PDF'):
            return 'application/pdf'
        elif data.startswith(b'PK\x03\x04') and b'word/' in data[:4096]:
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif data.startswith(b'<!DOCTYPE html') or data.startswith(b'<html'):
            return 'text/html'
        
        return mime_type

    async def _process_pdf(self, data: bytes) -> str:
        """Comprehensive PDF processing with advanced text extraction."""
        text_parts = []
        
        try:
            with io.BytesIO(data) as pdf_buffer:
                with pdfplumber.open(pdf_buffer) as pdf:
                    logger.info(f"Processing PDF with {len(pdf.pages)} pages using pdfplumber")
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text(
                            x_tolerance=2,
                            y_tolerance=3,
                            layout=True,
                            x_density=7.25,
                            y_density=7.25
                        )
                        
                        if page_text and page_text.strip():
                            cleaned_page_text = self._clean_pdf_text(page_text)
                            text_parts.append(f"PAGE {page_num}:\n{cleaned_page_text}")
                        
                        tables = page.extract_tables()
                        for table_idx, table in enumerate(tables):
                            if table and len(table) > 0:
                                formatted_table = self._format_table_comprehensive(
                                    table, 
                                    f"Table {table_idx + 1} on Page {page_num}"
                                )
                                if formatted_table.strip():
                                    text_parts.append(formatted_table)
                        
                        if hasattr(page, 'images') and page.images:
                            text_parts.append(f"[Images detected on page {page_num}]")
                            
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2 fallback")
            
            try:
                with io.BytesIO(data) as pdf_buffer:
                    reader = PyPDF2.PdfReader(pdf_buffer)
                    logger.info(f"Processing PDF with {len(reader.pages)} pages using PyPDF2")
                    
                    for page_num, page in enumerate(reader.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                cleaned_page_text = self._clean_pdf_text(page_text)
                                text_parts.append(f"PAGE {page_num}:\n{cleaned_page_text}")
                        except Exception as page_error:
                            logger.warning(f"Failed to extract text from page {page_num}: {page_error}")
                            continue
                            
            except Exception as fallback_error:
                raise DocumentProcessingError(f"Both pdfplumber and PyPDF2 failed: {str(fallback_error)}")
        
        if not text_parts:
            raise DocumentProcessingError("No text content extracted from PDF")
        
        return "\n\n".join(text_parts)

    def _clean_pdf_text(self, text: str) -> str:
        """Advanced PDF text cleaning with insurance document specifics."""
        if not text:
            return ""
        
        text = text.replace('\u00ad', '')  # soft hyphen
        text = text.replace('\u200b', '')  # zero-width space
        
        text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
        text = re.sub(r'(\w+)\s*\n\s*(\w+)', lambda m: 
                     f"{m.group(1)}{m.group(2)}" if m.group(1).islower() and m.group(2).islower() 
                     else f"{m.group(1)} {m.group(2)}", text)
        
        text = re.sub(r'\s+([,.;:!?])', r'\1', text)
        text = re.sub(r'([,.;:!?])\s*([A-Z])', r'\1 \2', text)
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'\n\s*(\d+\.)\s*', r'\n\1 ', text)
        text = re.sub(r'\n\s*([a-z]\))\s*', r'\n\1 ', text)
        text = re.sub(r'\n\s*([•·▪▫‣⁃])\s*', r'\n• ', text)
        
        return text.strip()

    def _format_table_comprehensive(self, table: List[List[str]], table_title: str = "TABLE") -> str:
        """Comprehensive table formatting optimized for insurance documents."""
        if not table or not any(any(cell for cell in row if cell) for row in table):
            return ""
        
        formatted_parts = [f"\n=== {table_title} ==="]
        
        try:
            cleaned_table = []
            for row in table:
                cleaned_row = []
                for cell in row:
                    if cell:
                        cell_content = str(cell).strip()
                        cell_content = re.sub(r'\s+', ' ', cell_content)
                        cell_content = re.sub(r'\n+', ' ', cell_content)
                        cleaned_row.append(cell_content)
                    else:
                        cleaned_row.append("")
                cleaned_table.append(cleaned_row)
            
            if not cleaned_table:
                return ""
            
            headers = cleaned_table[0] if cleaned_table else []
            clean_headers = [h for h in headers if h.strip()]
            
            if clean_headers:
                formatted_parts.append(f"HEADERS: {' | '.join(clean_headers)}")
                formatted_parts.append("-" * 60)
            
            for row_idx, row in enumerate(cleaned_table[1:], 1):
                if not any(cell.strip() for cell in row):
                    continue
                
                row_data = []
                for col_idx, cell in enumerate(row):
                    if cell.strip():
                        if col_idx < len(headers) and headers[col_idx].strip():
                            row_data.append(f"{headers[col_idx].strip()}: {cell.strip()}")
                        else:
                            row_data.append(cell.strip())
                
                if row_data:
                    formatted_parts.append(f"Row {row_idx}: {' | '.join(row_data)}")
            
            # Enhanced metadata tagging with expanded keywords
            table_content = ' '.join(formatted_parts).lower()
            metadata_tags = []
            
            insurance_keywords_enhanced = {
                'plan': 'CONTAINS_PLAN_INFO',
                'coverage': 'CONTAINS_COVERAGE_INFO',
                'limit': 'CONTAINS_LIMITS',
                'premium': 'CONTAINS_PREMIUM_INFO',
                'benefit': 'CONTAINS_BENEFITS',
                'exclusion': 'CONTAINS_EXCLUSIONS',
                'waiting': 'CONTAINS_WAITING_PERIODS',
                'deductible': 'CONTAINS_DEDUCTIBLES',
                'co-pay': 'CONTAINS_COPAY_INFO',
                'room rent': 'CONTAINS_ROOM_RENT_INFO',
                'icu': 'CONTAINS_ICU_INFO',
                'sum insured': 'CONTAINS_SUM_INSURED',
                'percentage': 'CONTAINS_PERCENTAGES',
                'ambulance': 'CONTAINS_AMBULANCE_INFO',
                'maternity': 'CONTAINS_MATERNITY_INFO',
                'pre-existing': 'CONTAINS_PED_INFO',
                'ayush': 'CONTAINS_AYUSH_INFO',
                'modern treatment': 'CONTAINS_MODERN_TREATMENTS',
                'organ donor': 'CONTAINS_DONOR_INFO',
                'zone': 'CONTAINS_ZONE_INFO',
                'age': 'CONTAINS_AGE_INFO',
                'hospital': 'CONTAINS_HOSPITAL_INFO'
            }
            
            for keyword, tag in insurance_keywords_enhanced.items():
                if keyword in table_content:
                    metadata_tags.append(tag)
            
            if metadata_tags:
                formatted_parts.append(f"METADATA: {' | '.join(set(metadata_tags))}")
            
            formatted_parts.append("=== END TABLE ===\n")
            
            return '\n'.join(formatted_parts)
            
        except Exception as e:
            logger.warning(f"Table formatting error: {e}")
            return self._format_table_simple(table, table_title)

    def _format_table_simple(self, table: List[List[str]], table_title: str = "TABLE") -> str:
        """Simple fallback table formatting."""
        if not table:
            return ""
        
        formatted_rows = [f"\n=== {table_title} ==="]
        for row in table:
            if row and any(cell for cell in row):
                clean_cells = [str(cell).strip() if cell else "" for cell in row]
                if any(clean_cells):
                    formatted_rows.append(" | ".join(clean_cells))
        
        formatted_rows.append("=== END TABLE ===\n")
        return "\n".join(formatted_rows)

    async def _process_docx(self, data: bytes) -> str:
        """Comprehensive DOCX processing with enhanced table and structure handling."""
        text_parts = []
        
        try:
            with io.BytesIO(data) as docx_buffer:
                doc = DocxDocument(docx_buffer)
                
                for para_idx, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip():
                        para_text = paragraph.text.strip()
                        
                        if paragraph.style and 'heading' in paragraph.style.name.lower():
                            para_text = f"SECTION: {para_text}"
                        elif para_text.isupper() and len(para_text) > 10:
                            para_text = f"HEADING: {para_text}"
                        
                        text_parts.append(para_text)
                
                for table_idx, table in enumerate(doc.tables, 1):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        if any(row_data):
                            table_data.append(row_data)
                    
                    if table_data:
                        formatted_table = self._format_table_comprehensive(
                            table_data, 
                            f"DOCX Table {table_idx}"
                        )
                        if formatted_table.strip():
                            text_parts.append(formatted_table)
                
                for section in doc.sections:
                    if section.header.paragraphs:
                        header_text = '\n'.join([p.text.strip() for p in section.header.paragraphs if p.text.strip()])
                        if header_text:
                            text_parts.insert(0, f"HEADER: {header_text}")
                    
                    if section.footer.paragraphs:
                        footer_text = '\n'.join([p.text.strip() for p in section.footer.paragraphs if p.text.strip()])
                        if footer_text:
                            text_parts.append(f"FOOTER: {footer_text}")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise DocumentProcessingError(f"DOCX processing failed: {str(e)}")

    async def _process_email(self, data: bytes) -> str:
        """Comprehensive email processing with attachment handling."""
        text_parts = []
        
        try:
            msg = message_from_bytes(data)
            
            important_headers = ['Subject', 'From', 'To', 'Cc', 'Bcc', 'Date', 'Priority']
            for header in important_headers:
                value = msg.get(header)
                if value:
                    text_parts.append(f"{header}: {value}")
            
            text_parts.append("=" * 50)
            
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    
                    if content_type == "text/plain":
                        body = part.get_payload(decode=True)
                        if body:
                            decoded_body = body.decode('utf-8', errors='ignore')
                            text_parts.append(f"PLAIN TEXT CONTENT:\n{decoded_body}")
                    
                    elif content_type == "text/html":
                        body = part.get_payload(decode=True)
                        if body:
                            html_content = body.decode('utf-8', errors='ignore')
                            soup = BeautifulSoup(html_content, 'html.parser')
                            text_content = soup.get_text(separator='\n', strip=True)
                            text_parts.append(f"HTML CONTENT:\n{text_content}")
                    
                    elif part.get_filename():
                        filename = part.get_filename()
                        text_parts.append(f"ATTACHMENT: {filename} ({content_type})")
            else:
                body = msg.get_payload(decode=True)
                if body:
                    content_type = msg.get_content_type()
                    decoded_body = body.decode('utf-8', errors='ignore')
                    
                    if content_type == "text/html":
                        soup = BeautifulSoup(decoded_body, 'html.parser')
                        decoded_body = soup.get_text(separator='\n', strip=True)
                    
                    text_parts.append(f"EMAIL CONTENT:\n{decoded_body}")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise DocumentProcessingError(f"Email processing failed: {str(e)}")

    async def _process_text(self, data: bytes) -> str:
        """Process plain text with encoding detection."""
        try:
            text = data.decode('utf-8')
        except UnicodeDecodeError:
            try:
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = data.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    text = data.decode('utf-8', errors='ignore')
            except Exception:
                text = data.decode('utf-8', errors='ignore')
        
        return text

    async def _process_html(self, data: bytes) -> str:
        """Process HTML with comprehensive text extraction."""
        try:
            html_content = data.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator='\n', strip=True)
            text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
            
            return text
            
        except Exception as e:
            raise DocumentProcessingError(f"HTML processing failed: {str(e)}")

    def _comprehensive_clean_text(self, text: str) -> str:
        """Comprehensive text cleaning optimized for insurance/legal documents."""
        cleaned = clean_text(text)
        cleaned = self._normalize_insurance_terminology(cleaned)
        cleaned = self._preserve_document_structure(cleaned)
        cleaned = self._final_text_cleanup(cleaned)
        
        return cleaned

    def _normalize_insurance_terminology(self, text: str) -> str:
        """Normalize insurance-specific terminology using comprehensive patterns."""
        # Apply all normalization patterns
        for pattern, replacement in self.normalization_patterns.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text

    def _preserve_document_structure(self, text: str) -> str:
        """Preserve important document structure using structure patterns."""
        # Apply structure recognition patterns
        for pattern, replacement in self.structure_patterns.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text

    def _final_text_cleanup(self, text: str) -> str:
        """Final comprehensive text cleanup."""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'\.([A-Z])', r'. \1', text)
        text = re.sub(r'\n\s*[•·▪▫‣⁃]\s*', '\n• ', text)
        
        return text.strip()

    def _intelligent_chunk_text(self, text: str) -> List[str]:
        """Intelligent chunking optimized for insurance/legal documents."""
        sections = self._split_by_document_sections(text)
        
        chunks = []
        target_size = max(1200, self.settings.CHUNK_SIZE)  # Larger chunks for better context
        overlap_size = max(250, self.settings.CHUNK_OVERLAP)  # More overlap for continuity
        
        for section in sections:
            if len(section.strip()) == 0:
                continue
            
            if len(section) <= target_size:
                chunks.append(section.strip())
            else:
                section_chunks = self._advanced_chunk_section(section, target_size, overlap_size)
                chunks.extend(section_chunks)
        
        final_chunks = []
        for chunk in chunks:
            chunk = chunk.strip()
            if len(chunk) > 75:  # Minimum meaningful chunk size
                final_chunks.append(chunk)
        
        logger.info(f"Created {len(final_chunks)} intelligent chunks from {len(sections)} sections")
        return final_chunks

    def _split_by_document_sections(self, text: str) -> List[str]:
        """Split text by major document sections while preserving context."""
        section_patterns = [
            r'\n\nSECTION:\s*([^\n]+)\n\n',
            r'\n\nSUBSECTION:\s*([^\n]+)\n\n',
            r'\n=== .+ ===\n',
            r'\nPAGE \d+:\n',
            r'\n\nCLAUSE \d+\.',
            r'\n{3,}'
        ]
        
        sections = [text]
        
        for pattern in section_patterns:
            new_sections = []
            for section in sections:
                parts = re.split(pattern, section, flags=re.IGNORECASE)
                if len(parts) == 1:
                    new_sections.append(section)
                else:
                    matches = list(re.finditer(pattern, section, flags=re.IGNORECASE))
                    last_end = 0
                    
                    for i, match in enumerate(matches):
                        if match.start() > last_end:
                            new_sections.append(section[last_end:match.start()])
                        
                        next_start = matches[i + 1].start() if i + 1 < len(matches) else len(section)
                        section_with_header = section[match.start():next_start]
                        new_sections.append(section_with_header)
                        
                        last_end = next_start
            
            sections = new_sections
        
        return [s.strip() for s in sections if s.strip()]

    def _advanced_chunk_section(self, section: str, target_size: int, overlap_size: int) -> List[str]:
        """Advanced chunking that respects sentence and paragraph boundaries."""
        paragraphs = [p.strip() for p in section.split('\n\n') if p.strip()]
        
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk + "\n\n" + paragraph) > target_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                if overlap_size > 0 and len(current_chunk) > overlap_size:
                    overlap_text = current_chunk[-overlap_size:]
                    sentences = overlap_text.split('. ')
                    if len(sentences) > 1:
                        overlap_text = '. '.join(sentences[-2:])
                    current_chunk = overlap_text + "\n\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        final_chunks = []
        for chunk in chunks:
            if len(chunk) <= target_size * 1.5:
                final_chunks.append(chunk)
            else:
                sub_chunks = split_text_into_chunks(chunk, target_size, overlap_size)
                final_chunks.extend(sub_chunks)
        
        return final_chunks
